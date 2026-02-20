"""Step 7: Final Output Generation (generator.py)

Generates the complete resume package:
1. Resume PDF (reportlab, pixel-perfect typography)
2. Resume DOCX (python-docx, closest approximation)
3. score_report.json
4. keyword_coverage.json
5. reframing_log.json
6. interview_prep.md
7. iteration_log.json
8. format_warnings.json

Output saved to: output/{company_name}_{date}/
"""

import json
import logging
import os
import re
from datetime import datetime

from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

logger = logging.getLogger(__name__)


class QualityGateBlockedError(Exception):
    """Raised when critical rule failures block PDF generation."""
    def __init__(self, message: str, blocked_reason: str, rule13_failures: list):
        super().__init__(message)
        self.blocked_reason = blocked_reason
        self.rule13_failures = rule13_failures


import re as _re


def _fix_sp(t):
    """Nuclear spacing fix — applied to every text before PDF render."""
    if not t:
        return t
    t = _re.sub(r'(\d)([a-z])', r'\1 \2', t)
    t = _re.sub(r'(\d\.?\d*)([A-Z][a-z]{2,})', r'\1 \2', t)
    t = _re.sub(r'(%)([a-zA-Z])', r'\1 \2', t)  # "35%improvement" -> "35% improvement"
    t = _re.sub(r'  +', ' ', t)
    return t


# --- Page geometry ---
PAGE_W, PAGE_H = A4  # 595.9 x 842.9
MARGIN_LR = 0.54 * inch  # 38.7pt
MARGIN_TOP = 0.5 * inch
MARGIN_BOTTOM = 0.5 * inch
CONTENT_W = PAGE_W - 2 * MARGIN_LR

# --- Colors ---
BLACK = HexColor("#000000")
GREEN = HexColor("#1CAD62")
GRAY = HexColor("#3E3E3E")
WHITE = HexColor("#FFFFFF")

# --- Font registration ---
_FONTS_REGISTERED = False


def _register_fonts():
    """Register Georgia and Arial/Helvetica as fallback fonts.

    Georgia (serif) replaces Volkhov; Helvetica (built-in) replaces PT Sans.
    If Georgia TTFs are available on the system, register them.
    Otherwise fall back to reportlab built-ins (Times-Roman, Helvetica).
    """
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return

    # Helvetica and Times-Roman are built into reportlab — always available
    # Try to register Georgia if available on macOS
    georgia_paths = [
        "/Library/Fonts/Georgia.ttf",
        "/Library/Fonts/Georgia Bold.ttf",
        "/System/Library/Fonts/Supplemental/Georgia.ttf",
        "/System/Library/Fonts/Supplemental/Georgia Bold.ttf",
    ]
    arial_paths = [
        "/Library/Fonts/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    ]

    # Try Georgia (serif — for name and section headers)
    georgia_reg = None
    georgia_bold = None
    for p in georgia_paths:
        if os.path.exists(p) and "Bold" not in p and georgia_reg is None:
            georgia_reg = p
        if os.path.exists(p) and "Bold" in p and georgia_bold is None:
            georgia_bold = p

    if georgia_reg:
        try:
            pdfmetrics.registerFont(TTFont("Georgia", georgia_reg))
            if georgia_bold:
                pdfmetrics.registerFont(TTFont("Georgia-Bold", georgia_bold))
            else:
                pdfmetrics.registerFont(TTFont("Georgia-Bold", georgia_reg))
        except Exception:
            georgia_reg = None

    # Try Arial (sans-serif — for body text)
    arial_reg = None
    arial_bold = None
    for p in arial_paths:
        if os.path.exists(p) and "Bold" not in p and arial_reg is None:
            arial_reg = p
        if os.path.exists(p) and "Bold" in p and arial_bold is None:
            arial_bold = p

    if arial_reg:
        try:
            pdfmetrics.registerFont(TTFont("Arial", arial_reg))
            if arial_bold:
                pdfmetrics.registerFont(TTFont("Arial-Bold", arial_bold))
            else:
                pdfmetrics.registerFont(TTFont("Arial-Bold", arial_reg))
        except Exception:
            arial_reg = None

    _FONTS_REGISTERED = True

    # Set globals for font names to use
    global SERIF_FONT, SERIF_FONT_BOLD, SANS_FONT, SANS_FONT_BOLD
    SERIF_FONT = "Georgia" if georgia_reg else "Times-Roman"
    SERIF_FONT_BOLD = "Georgia-Bold" if georgia_reg and georgia_bold else (
        "Georgia" if georgia_reg else "Times-Bold"
    )
    SANS_FONT = "Arial" if arial_reg else "Helvetica"
    SANS_FONT_BOLD = "Arial-Bold" if arial_reg and arial_bold else (
        "Arial" if arial_reg else "Helvetica-Bold"
    )


# Defaults (overridden by _register_fonts)
SERIF_FONT = "Times-Roman"
SERIF_FONT_BOLD = "Times-Bold"
SANS_FONT = "Helvetica"
SANS_FONT_BOLD = "Helvetica-Bold"

# --- Font sizes (in points) ---
NAME_SIZE = 17.8
SUBTITLE_SIZE = 12.7
CONTACT_SIZE = 8.9
SECTION_HEADER_SIZE = 12.7
COMPANY_SIZE = 12.7
LOCATION_SIZE = 10.1
TITLE_SIZE = 10.1
DATES_SIZE = 10.1
BULLET_SIZE = 9.5
SKILLS_SIZE = 9.5
CERT_SIZE = 10.1
CERT_ISSUER_SIZE = 9.5

# --- Spacing (ATS-optimized: tighter header for cleaner scan) ---
SPACE_AFTER_NAME = 1
SPACE_AFTER_SUBTITLE = 1
SPACE_AFTER_CONTACT = 8
SPACE_AFTER_SECTION_RULE = 8
SPACE_BETWEEN_ROLES = 10
SPACE_BETWEEN_BULLETS = 1
SPACE_BEFORE_SECTION = 12
RULE_HEIGHT = 0.6
BULLET_INDENT = 11.3
BULLET_CHAR = "\u2022"


def _dedup_edu_location(institution: str, location: str) -> str:
    """Strip city from education location if institution already ends with it.

    E.g. institution="Indian Institute of Management, Shillong", location="Shillong, India"
    -> returns "India" (not "Shillong, India")
    """
    if not location or not institution:
        return location
    inst_parts = [p.strip() for p in institution.split(",")]
    if len(inst_parts) >= 2:
        inst_city = inst_parts[-1].strip().lower()
        loc_parts = [p.strip() for p in location.split(",")]
        if loc_parts and loc_parts[0].strip().lower() == inst_city:
            # Remove the duplicated city, keep only country
            remaining = ", ".join(loc_parts[1:]).strip()
            return remaining if remaining else location
    return location


def _build_degree_with_field(edu: dict) -> str:
    """Build degree string, appending field/specialization if available.

    E.g. degree="Bachelor of Technology (BTech)", field="Computer Science"
    -> "Bachelor of Technology (BTech), Computer Science"
    """
    degree = (edu.get("degree") or "").strip()
    field = (edu.get("field") or "").strip()
    if field and field.lower() not in degree.lower():
        degree = f"{degree}, {field}" if degree else field
    return degree


def _esc(text: str) -> str:
    """Escape text for reportlab Paragraph XML."""
    if not text:
        return ""
    # Spacing fix for all text
    text = re.sub(r'(\d)([a-z])', r'\1 \2', text)
    text = re.sub(r'(\d\.?\d*)([A-Z][a-z]{2,})', r'\1 \2', text)
    text = re.sub(r'(%)([a-zA-Z])', r'\1 \2', text)  # "35%improvement" -> "35% improvement"
    text = re.sub(r'  +', ' ', text)
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


ACRONYMS_TO_BOLD = ("B2B", "B2C", "SMB", "GTM", "FP&A", "OKR", "KPI", "ROI", "API", "CRM")


def _bold_metrics(text: str, font_name: str, bold_font_name: str,
                  font_size: float, color_hex: str) -> str:
    """Convert bullet text to Paragraph XML with bold numbers/percentages/metrics.

    Bolds: numbers with %, $, x/×, standalone numbers >= 2 digits,
    short metric phrases like '2.5× adoption', and acronyms (B2B, B2C, SMB, etc.).
    """
    # Nuclear spacing fix — runs on every text before PDF render
    if text:
        text = re.sub(r'(\d)([a-z])', r'\1 \2', text)
        text = re.sub(r'(\d\.?\d*)([A-Z][a-z]{2,})', r'\1 \2', text)
        text = re.sub(r'(%)([a-zA-Z])', r'\1 \2', text)  # "35%improvement" -> "35% improvement"
        text = re.sub(r'([a-zA-Z0-9])\(', r'\1 (', text)
        text = re.sub(r'  +', ' ', text)
    if not text:
        return ""
    # Step 1: Replace acronyms with placeholders (no digits) so metric pattern won't touch them
    # Use digit-free placeholders for B2B/B2C so metric regex doesn't match "2B"/"2C"
    _ACR_PLACEHOLDERS = {"B2B": "\uE000BtoB\uE001", "B2C": "\uE000BtoC\uE001"}
    placeholders = {}
    for acr in ACRONYMS_TO_BOLD:
        if acr in text:
            ph = _ACR_PLACEHOLDERS.get(acr) or f"\uE000{acr.replace('&', '_').replace('2', 'to')}\uE001"
            placeholders[ph] = acr
            text = re.sub(re.escape(acr), ph, text)
    # Pattern: number with optional decimal + optional suffix (%, x, ×, +, K, M, B)
    # Also catches dollar amounts and ranges
    metric_pattern = re.compile(
        r'(\$?\d[\d,]*\.?\d*\s*(?:[%×x+]|[KMB]\b)?'  # core number
        r'(?:\s*[-–]\s*\$?\d[\d,]*\.?\d*\s*(?:[%×x+]|[KMB]\b)?)?'  # optional range
        r'(?:\s+(?:months?|years?|users?|partners?|clients?|agents?|markets?|businesses))?)'  # optional unit
    )

    parts = []
    last_end = 0
    for m in metric_pattern.finditer(text):
        start, end = m.span()
        if start > last_end:
            parts.append(_esc(text[last_end:start]))
        matched = m.group(0).strip()
        if matched and any(c.isdigit() for c in matched):
            parts.append(
                f'<b><font name="{bold_font_name}" size="{font_size}">'
                f'{_esc(matched)}</font></b>'
            )
        else:
            parts.append(_esc(matched))
        # Use start+len(matched) so trailing space (e.g. "8 " -> "8") is preserved in next part
        last_end = start + len(matched)
    if last_end < len(text):
        parts.append(_esc(text[last_end:]))
    # ReportLab Paragraph collapses spaces between inline elements. Replace space
    # after </font></b> with &#160; (non-breaking space) so "35% improvement" etc. render correctly.
    result = "".join(parts)
    result = re.sub(r'(</font></b>)\s+', r'\1&#160;', result)
    # Step 2: Restore acronym placeholders as bold
    for ph, acr in placeholders.items():
        bold_acr = f'<b><font name="{bold_font_name}" size="{font_size}">{_esc(acr)}</font></b>'
        result = result.replace(_esc(ph), bold_acr)
    return result


def _build_styles():
    """Build all ParagraphStyles used in the resume."""
    _register_fonts()
    styles = {}
    styles["name"] = ParagraphStyle(
        "name", fontName=SERIF_FONT_BOLD, fontSize=NAME_SIZE,
        textColor=BLACK, alignment=TA_CENTER, leading=NAME_SIZE + 4,
        spaceAfter=SPACE_AFTER_NAME,
    )
    styles["subtitle"] = ParagraphStyle(
        "subtitle", fontName=SANS_FONT, fontSize=SUBTITLE_SIZE,
        textColor=GREEN, alignment=TA_CENTER, leading=SUBTITLE_SIZE + 4,
        spaceAfter=SPACE_AFTER_SUBTITLE,
    )
    styles["contact"] = ParagraphStyle(
        "contact", fontName=SANS_FONT, fontSize=CONTACT_SIZE,
        textColor=GRAY, alignment=TA_CENTER, leading=CONTACT_SIZE + 4,
        spaceAfter=SPACE_AFTER_CONTACT,
    )
    styles["section_header"] = ParagraphStyle(
        "section_header", fontName=SERIF_FONT, fontSize=SECTION_HEADER_SIZE,
        textColor=BLACK, alignment=TA_CENTER, leading=SECTION_HEADER_SIZE + 4,
        spaceAfter=2,
    )
    styles["company"] = ParagraphStyle(
        "company", fontName=SANS_FONT, fontSize=COMPANY_SIZE,
        textColor=GREEN, alignment=TA_LEFT, leading=COMPANY_SIZE + 4,
    )
    styles["location"] = ParagraphStyle(
        "location", fontName=SANS_FONT, fontSize=LOCATION_SIZE,
        textColor=GRAY, alignment=TA_RIGHT, leading=LOCATION_SIZE + 3,
    )
    styles["job_title"] = ParagraphStyle(
        "job_title", fontName=SANS_FONT, fontSize=TITLE_SIZE,
        textColor=BLACK, alignment=TA_LEFT, leading=TITLE_SIZE + 3,
    )
    styles["dates"] = ParagraphStyle(
        "dates", fontName=SANS_FONT, fontSize=DATES_SIZE,
        textColor=GRAY, alignment=TA_RIGHT, leading=DATES_SIZE + 3,
    )
    styles["bullet"] = ParagraphStyle(
        "bullet", fontName=SANS_FONT, fontSize=BULLET_SIZE,
        textColor=GRAY, alignment=TA_JUSTIFY, leading=BULLET_SIZE + 3,
        leftIndent=BULLET_INDENT, firstLineIndent=-BULLET_INDENT + 2.5,
        spaceBefore=0, spaceAfter=SPACE_BETWEEN_BULLETS,
        allowWidows=0,
    )
    styles["skills"] = ParagraphStyle(
        "skills", fontName=SANS_FONT, fontSize=SKILLS_SIZE,
        textColor=GRAY, alignment=TA_LEFT, leading=SKILLS_SIZE + 5,
    )
    styles["summary"] = ParagraphStyle(
        "summary", fontName=SANS_FONT, fontSize=BULLET_SIZE,
        textColor=GRAY, alignment=TA_JUSTIFY, leading=BULLET_SIZE + 4.5,
        allowWidows=0,
    )
    styles["edu_institution"] = ParagraphStyle(
        "edu_institution", fontName=SANS_FONT, fontSize=COMPANY_SIZE,
        textColor=GREEN, alignment=TA_LEFT, leading=COMPANY_SIZE + 4,
    )
    styles["edu_detail"] = ParagraphStyle(
        "edu_detail", fontName=SANS_FONT, fontSize=TITLE_SIZE,
        textColor=GRAY, alignment=TA_LEFT, leading=TITLE_SIZE + 3,
    )
    styles["cert"] = ParagraphStyle(
        "cert", fontName=SANS_FONT, fontSize=CERT_SIZE,
        textColor=GRAY, alignment=TA_LEFT, leading=CERT_SIZE + 4,
    )
    styles["award_title"] = ParagraphStyle(
        "award_title", fontName=SANS_FONT_BOLD, fontSize=TITLE_SIZE,
        textColor=GRAY, alignment=TA_LEFT, leading=TITLE_SIZE + 3,
    )
    styles["award_desc"] = ParagraphStyle(
        "award_desc", fontName=SANS_FONT, fontSize=BULLET_SIZE,
        textColor=GRAY, alignment=TA_LEFT, leading=BULLET_SIZE + 3,
    )
    return styles


class _HRLine:
    """Flowable that draws a thin horizontal rule spanning content width."""

    def __init__(self, width, color=BLACK, thickness=RULE_HEIGHT):
        self.width = width
        self.color = color
        self.thickness = thickness
        self.height = thickness + 4  # padding below
        self.spaceBefore = 0
        self.spaceAfter = SPACE_AFTER_SECTION_RULE

    def wrap(self, available_width, available_height):
        return (self.width, self.height)

    def draw(self):
        pass

    def drawOn(self, canvas, x, y, _sW=0):
        canvas.saveState()
        canvas.setStrokeColor(self.color)
        canvas.setLineWidth(self.thickness)
        canvas.line(x, y + self.height - 1, x + self.width, y + self.height - 1)
        canvas.restoreState()


from reportlab.platypus import Flowable


class HRLineFlowable(Flowable):
    """Proper Flowable subclass for horizontal rule."""

    def __init__(self, width, color=BLACK, thickness=RULE_HEIGHT):
        super().__init__()
        self.width = width
        self.color = color
        self.thickness = thickness
        self.spaceAfter = SPACE_AFTER_SECTION_RULE

    def wrap(self, available_width, available_height):
        return (self.width, self.thickness + 2)

    def draw(self):
        self.canv.saveState()
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 1, self.width, 1)
        self.canv.restoreState()


def _two_col_table(left_text, left_style, right_text, right_style, col_widths=None):
    """Create a two-column table for company/location or title/dates rows."""
    if col_widths is None:
        col_widths = [CONTENT_W * 0.7, CONTENT_W * 0.3]
    left_para = Paragraph(_fix_sp(left_text), left_style)
    right_para = Paragraph(_fix_sp(right_text), right_style)
    t = Table([[left_para, right_para]], colWidths=col_widths)
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
    ]))
    return t


def _clean_url(url: str) -> str:
    """Remove https:// and trailing slash for display: 'https://linkedin.com/in/foo/' -> 'linkedin.com/in/foo'."""
    if not url:
        return ""
    url = re.sub(r'^https?://', '', url)
    return url.rstrip("/")


def _ensure_url(url: str) -> str:
    """Ensure URL has scheme for href. Adds https:// if missing."""
    if not url or not isinstance(url, str):
        return ""
    url = (url or "").strip()
    if not url:
        return ""
    if not url.startswith(("http://", "https://")):
        return "https://" + url
    return url


def _escape_href(url: str) -> str:
    """Escape URL for safe use in href attribute (XML/HTML)."""
    if not url:
        return ""
    return (url.replace("&", "&amp;").replace('"', "&quot;")
            .replace("<", "&lt;").replace(">", "&gt;"))


def _link_markup(url: str, label: str, color: str = None) -> str:
    """Build reportlab <a href="...">label</a> markup for clickable link. ATS-compliant."""
    if not url:
        return ""
    href = _escape_href(_ensure_url(url))
    if color:
        return f'<a href="{href}" color="{color}">{_esc(label)}</a>'
    return f'<a href="{href}">{_esc(label)}</a>'


def _clean_spacing(text: str) -> str:
    """Collapse multiple spaces into one and strip leading/trailing whitespace."""
    if not text:
        return text
    return re.sub(r'  +', ' ', text).strip()


def _fix_acronym_casing(text: str) -> str:
    """Prevent .title() from turning AI->Ai, ML->Ml, etc. Preserve common acronyms."""
    if not text:
        return text
    for acr in ("AI", "ML", "API", "CRM", "GTM", "FP&A", "B2B", "B2C"):
        text = re.sub(r'\b' + re.escape(acr) + r'\b', acr, text, flags=re.IGNORECASE)
    return text


def _build_title_display(title: str, role_desc: str) -> str:
    """Combine title and role_description, deduplicating overlapping terms.

    Handles cases like:
      title="Senior Product Manager - AI", role_desc="AI-enabled customer success platform"
      -> "Senior Product Manager – AI-enabled customer success platform"
    """
    if not role_desc:
        return title

    # Extract the suffix after the last " - " in the title (e.g. "AI" from "Senior PM - AI")
    title_suffix = ""
    if " - " in title:
        title_suffix = title.rsplit(" - ", 1)[1].strip()

    if title_suffix:
        # Check if role_desc starts with the same term (case-insensitive)
        rd_first_word = role_desc.split("-")[0].split()[0] if role_desc else ""
        if title_suffix.lower() == rd_first_word.lower():
            # Remove the suffix from title, use role_desc as-is
            base_title = title.rsplit(" - ", 1)[0].strip()
            return f"{base_title} – {role_desc}"

    return f"{title} – {role_desc}"


def _format_cert(cert) -> str:
    """Format a certification entry to 'Name — Issuer (Year)' string."""
    if isinstance(cert, str):
        return cert
    if isinstance(cert, dict):
        name = cert.get("name", "")
        issuer = cert.get("issuer", "")
        year = cert.get("year", "")
        parts = [name]
        if issuer:
            parts.append(f"— {issuer}")
        if year:
            parts.append(f"({year})")
        return " ".join(parts) if name else str(cert)
    return str(cert)


def _generate_pdf(content: dict, pkb: dict, output_path: str):
    """Generate the pixel-perfect resume PDF using reportlab."""
    _register_fonts()
    styles = _build_styles()

    # Personal info from PKB or content
    personal = pkb.get("personal_info") or {}
    name = personal.get("name") or "Candidate"
    email = personal.get("email") or ""
    phone = personal.get("phone") or ""
    linkedin = personal.get("linkedin_url") or ""
    location = personal.get("location") or ""
    # Normalize location to City, Country
    if location:
        parts = [p.strip() for p in location.split(",")]
        if len(parts) >= 2:
            location = f"{parts[0]}, {parts[-1]}"

    story = []

    # --- HEADER ---
    story.append(Paragraph(_fix_sp(_esc(name)), styles["name"]))

    # Subtitle: use pre-generated tagline from reframer, or fallback
    subtitle = _clean_spacing(content.get("subtitle") or "Senior Product Manager | Enterprise Products | 8+ Years")
    summary_text = _clean_spacing((content.get("professional_summary") or "").strip())
    summary_lines = [l.strip() for l in summary_text.split("\n") if l.strip()]

    story.append(Paragraph(_fix_sp(_esc(subtitle)), styles["subtitle"]))

    # Contact: single line — phone, email, LinkedIn, GitHub, Portfolio, location (ATS-friendly)
    github = personal.get("github_url") or ""
    portfolio = personal.get("portfolio_url") or ""
    contact_parts = []
    if phone:
        contact_parts.append(_esc(phone))
    if email:
        contact_parts.append(_esc(email))
    if linkedin:
        contact_parts.append(_link_markup(linkedin, "LinkedIn", color="#0077B5"))
    if github:
        contact_parts.append(_link_markup(github, "Github", color="#0066CC"))
    if portfolio:
        contact_parts.append(_link_markup(portfolio, "Portfolio", color="#0066CC"))
    if location:
        contact_parts.append(_esc(location))
    if contact_parts:
        contact_line = " \u2022 ".join(contact_parts)
        story.append(Paragraph(_fix_sp(contact_line), styles["contact"]))

    # --- SUMMARY SECTION ---
    story.append(Spacer(1, SPACE_BEFORE_SECTION))
    story.append(Paragraph(_fix_sp("Professional Summary"), styles["section_header"]))
    story.append(HRLineFlowable(CONTENT_W))
    story.append(Spacer(1, 2))

    # Render summary with bold metrics
    if summary_lines:
        full_summary = _clean_spacing(" ".join(summary_lines))
        summary_xml = _bold_metrics(
            full_summary, SANS_FONT, SANS_FONT_BOLD, BULLET_SIZE, "#3E3E3E"
        )
        story.append(Paragraph(_fix_sp(summary_xml), styles["summary"]))

    # --- EXPERIENCE SECTION ---
    work = content.get("work_experience") or []
    if work:
        story.append(Spacer(1, SPACE_BEFORE_SECTION))
        story.append(Paragraph(_fix_sp("Work Experience"), styles["section_header"]))
        story.append(HRLineFlowable(CONTENT_W))
        story.append(Spacer(1, 2))

        for idx, role in enumerate(work):
            company = role.get("company") or ""
            loc = role.get("location") or ""
            title = role.get("title") or ""
            dates = role.get("dates") or ""
            if isinstance(dates, dict):
                dates = f"{dates.get('start', '')} \u2013 {dates.get('end', '')}"

            # Build role elements, then wrap in KeepTogether for widow/orphan control
            role_elements = []
            if idx > 0:
                role_elements.append(Spacer(1, SPACE_BETWEEN_ROLES))

            # Company + Location row
            company_xml = f'<font name="{SANS_FONT}" size="{COMPANY_SIZE}" color="#1CAD62">{_esc(company)}</font>'
            loc_xml = f'<font name="{SANS_FONT}" size="{LOCATION_SIZE}" color="#3E3E3E">{_esc(loc)}</font>'
            role_elements.append(_two_col_table(company_xml, styles["company"], loc_xml, styles["location"]))

            # Title + Dates row (with optional role_description: "Title – AI enabled platform for FP&A")
            role_desc = _fix_acronym_casing((role.get("role_description") or "").strip())
            title_display = _build_title_display(title, role_desc)
            title_xml = f'<font name="{SANS_FONT}" size="{TITLE_SIZE}" color="#000000">{_esc(title_display)}</font>'
            dates_xml = f'<font name="{SANS_FONT}" size="{DATES_SIZE}" color="#3E3E3E">{_esc(dates)}</font>'
            role_elements.append(_two_col_table(title_xml, styles["job_title"], dates_xml, styles["dates"]))

            # First 2 bullets kept with header (widow/orphan prevention)
            bullets = role.get("bullets") or []
            for bi in range(min(2, len(bullets))):
                bt = _clean_spacing(bullets[bi])
                bullet_xml = _bold_metrics(
                    f"{BULLET_CHAR}  {bt}",
                    SANS_FONT, SANS_FONT_BOLD, BULLET_SIZE, "#3E3E3E",
                )
                role_elements.append(Paragraph(_fix_sp(bullet_xml), styles["bullet"]))

            # KeepTogether: header + first 2 bullets stay on same page
            story.append(KeepTogether(role_elements))

            # Remaining bullets
            for bullet_text in bullets[2:]:
                bt = _clean_spacing(bullet_text)
                bullet_xml = _bold_metrics(
                    f"{BULLET_CHAR}  {bt}",
                    SANS_FONT, SANS_FONT_BOLD, BULLET_SIZE, "#3E3E3E",
                )
                story.append(Paragraph(_fix_sp(bullet_xml), styles["bullet"]))

    # --- KEY PROJECTS SECTION ---
    key_projects = content.get("key_projects") or []
    if key_projects:
        project_header_elements = [
            Spacer(1, SPACE_BEFORE_SECTION),
            Paragraph(_fix_sp("Key Projects"), styles["section_header"]),
            HRLineFlowable(CONTENT_W),
            Spacer(1, 2),
        ]

        for proj in key_projects:
            proj_name = (proj.get("name") or "").strip()
            proj_desc = (proj.get("description") or "").strip()

            # Project name (green, same as company name style)
            name_xml = f'<font name="{SANS_FONT}" size="{COMPANY_SIZE}" color="#1CAD62">{_esc(proj_name)}</font>'
            project_header_elements.append(Paragraph(_fix_sp(name_xml), styles["company"]))

            # Description (same style as role title)
            if proj_desc:
                desc_xml = f'<font name="{SANS_FONT}" size="{TITLE_SIZE}" color="#000000">{_esc(proj_desc)}</font>'
                project_header_elements.append(Paragraph(_fix_sp(desc_xml), styles["job_title"]))

            # Bullets
            for bullet_text in (proj.get("bullets") or []):
                bt = _clean_spacing(bullet_text)
                bullet_xml = _bold_metrics(
                    f"{BULLET_CHAR}  {bt}",
                    SANS_FONT, SANS_FONT_BOLD, BULLET_SIZE, "#3E3E3E",
                )
                project_header_elements.append(Paragraph(_fix_sp(bullet_xml), styles["bullet"]))

        # KeepTogether: header + first project stay on same page
        story.append(KeepTogether(project_header_elements))

    # --- SKILLS SECTION ---
    skills = content.get("skills") or {}
    all_skills = []
    for category in ("technical", "methodologies", "domains"):
        all_skills.extend(skills.get(category) or [])
    if all_skills:
        # Build skill paragraphs
        skill_parts = []
        for cat_name, cat_key in [("Technical", "technical"), ("Methodologies", "methodologies"), ("Domains", "domains")]:
            items = skills.get(cat_key) or []
            if items:
                label = f'<b><font name="{SANS_FONT_BOLD}" size="{SKILLS_SIZE}">{cat_name}:</font></b> '
                skill_parts.append(label + _esc(_clean_spacing("  |  ".join(items))))

        # Wrap header + HR + first category in KeepTogether to prevent page-break split
        skills_header_elements = [
            Spacer(1, SPACE_BEFORE_SECTION),
            Paragraph(_fix_sp("Skills"), styles["section_header"]),
            HRLineFlowable(CONTENT_W),
            Spacer(1, 2),
        ]
        if skill_parts:
            skills_header_elements.append(Paragraph(_fix_sp(skill_parts[0]), styles["skills"]))
            skills_header_elements.append(Spacer(1, 2))
        story.append(KeepTogether(skills_header_elements))

        # Remaining skill categories
        for part in skill_parts[1:]:
            story.append(Paragraph(_fix_sp(part), styles["skills"]))
            story.append(Spacer(1, 2))

    # --- AWARDS SECTION ---
    awards = content.get("awards") or []
    if awards:
        story.append(Spacer(1, SPACE_BEFORE_SECTION))
        story.append(Paragraph(_fix_sp("Awards &amp; Recognition"), styles["section_header"]))
        story.append(HRLineFlowable(CONTENT_W))
        story.append(Spacer(1, 2))

        # Always render awards as simple bullet list (Part B Fix 6)
        for award in awards:
            clean = _clean_spacing((award or "").strip().lstrip("\u2022").lstrip("•").strip())
            story.append(Paragraph(_fix_sp(f'{BULLET_CHAR}  {_esc(clean)}'), styles["bullet"]))

    # --- EDUCATION SECTION ---
    education = content.get("education") or []
    if education:
        story.append(Spacer(1, SPACE_BEFORE_SECTION))
        story.append(Paragraph(_fix_sp("Education"), styles["section_header"]))
        story.append(HRLineFlowable(CONTENT_W))
        story.append(Spacer(1, 2))

        for edu in education:
            institution = edu.get("institution") or ""
            degree = _build_degree_with_field(edu)
            dates = edu.get("dates") or ""
            if isinstance(dates, dict):
                dates = f"{dates.get('start', '')} \u2013 {dates.get('end', '')}"
            edu_loc = _dedup_edu_location(institution, edu.get("location") or "")

            # Institution + Location row
            inst_xml = f'<font name="{SANS_FONT}" size="{COMPANY_SIZE}" color="#1CAD62">{_esc(institution)}</font>'
            loc_xml = f'<font name="{SANS_FONT}" size="{LOCATION_SIZE}" color="#3E3E3E">{_esc(edu_loc)}</font>'
            story.append(_two_col_table(inst_xml, styles["edu_institution"], loc_xml, styles["location"]))

            # Degree + Dates row
            deg_xml = f'<font name="{SANS_FONT}" size="{TITLE_SIZE}" color="#3E3E3E">{_esc(degree)}</font>'
            dates_xml = f'<font name="{SANS_FONT}" size="{DATES_SIZE}" color="#3E3E3E">{_esc(dates)}</font>'
            story.append(_two_col_table(deg_xml, styles["edu_detail"], dates_xml, styles["dates"]))
            story.append(Spacer(1, 4))

    # --- CERTIFICATIONS SECTION ---
    certs = content.get("certifications") or []
    if certs:
        story.append(Spacer(1, SPACE_BEFORE_SECTION))
        story.append(Paragraph(_fix_sp("Certifications"), styles["section_header"]))
        story.append(HRLineFlowable(CONTENT_W))
        story.append(Spacer(1, 2))

        # Render all certs on a single line separated by pipes (RULE 11)
        cert_texts = [_format_cert(c) for c in certs]
        joined = "  |  ".join(cert_texts)
        story.append(Paragraph(
            _fix_sp(_esc(joined)), styles["cert"]
        ))

    # --- Build PDF ---
    frame = Frame(
        MARGIN_LR, MARGIN_BOTTOM,
        CONTENT_W, PAGE_H - MARGIN_TOP - MARGIN_BOTTOM,
        leftPadding=0, rightPadding=0,
        topPadding=0, bottomPadding=0,
    )
    template = PageTemplate(id="resume", frames=[frame])
    doc = BaseDocTemplate(
        output_path, pagesize=A4,
        leftMargin=MARGIN_LR, rightMargin=MARGIN_LR,
        topMargin=MARGIN_TOP, bottomMargin=MARGIN_BOTTOM,
    )
    doc.addPageTemplates([template])

    # SPACING VERIFICATION — will print to console if any spacing bugs remain
    for element in story:
        if hasattr(element, 'text'):
            txt = str(element.text) if element.text else ""
            bugs = re.findall(r'\d[a-z]', txt)
            if bugs:
                print(f"⚠️ SPACING BUG IN PDF: '{txt[:80]}' — found: {bugs}")

    doc.build(story)
    logger.info("PDF generated: %s", output_path)


def _add_docx_hyperlink(paragraph, text: str, url: str, font_size=None, color_rgb=None):
    """Add a hyperlink run to a docx paragraph. Returns the hyperlink element."""
    from docx.oxml.parser import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    url = _ensure_url(url)
    if not url:
        return
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font_size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size)))  # half-points
        rPr.append(sz)
    if color_rgb is not None:
        color = OxmlElement("w:color")
        r, g, b = color_rgb[0], color_rgb[1], color_rgb[2]
        color.set(qn("w:val"), f"{r:02X}{g:02X}{b:02X}")
        rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _generate_docx(content: dict, pkb: dict, output_path: str):
    """Generate a DOCX version of the resume using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.54)
    section.right_margin = Inches(0.54)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    green_rgb = RGBColor(28, 173, 98)
    gray_rgb = RGBColor(62, 62, 62)
    black_rgb = RGBColor(0, 0, 0)

    personal = pkb.get("personal_info") or {}
    name = personal.get("name") or "Candidate"

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(17.8)
    run.font.color.rgb = black_rgb
    run.font.name = "Georgia"

    # Subtitle (use generated tagline from content)
    subtitle = content.get("subtitle") or "Senior Product Manager | Enterprise Products | 8+ Years"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12.7)
    run.font.color.rgb = green_rgb
    run.font.name = "Arial"

    # Contact: phone, email, LinkedIn (clickable), location; then Github, Portfolio (clickable)
    sep = " \u2022 "
    font_sz = 18  # 8.9pt in half-points
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts_added = []
    if personal.get("phone"):
        run = p.add_run(personal["phone"])
        run.font.size = Pt(8.9)
        run.font.color.rgb = gray_rgb
        run.font.name = "Arial"
        parts_added.append(True)
    if personal.get("email"):
        if parts_added:
            p.add_run(sep).font.size = Pt(8.9)
        run = p.add_run(personal["email"])
        run.font.size = Pt(8.9)
        run.font.color.rgb = gray_rgb
        run.font.name = "Arial"
        parts_added.append(True)
    if personal.get("linkedin_url"):
        if parts_added:
            p.add_run(sep).font.size = Pt(8.9)
        _add_docx_hyperlink(p, "LinkedIn", personal["linkedin_url"], font_size=font_sz, color_rgb=(0, 119, 181))
        parts_added.append(True)
    if personal.get("github_url"):
        if parts_added:
            p.add_run(sep).font.size = Pt(8.9)
        _add_docx_hyperlink(p, "Github", personal["github_url"], font_size=font_sz, color_rgb=(0, 102, 204))
        parts_added.append(True)
    if personal.get("portfolio_url"):
        if parts_added:
            p.add_run(sep).font.size = Pt(8.9)
        _add_docx_hyperlink(p, "Portfolio", personal["portfolio_url"], font_size=font_sz, color_rgb=(0, 102, 204))
        parts_added.append(True)
    if personal.get("location"):
        if parts_added:
            p.add_run(sep).font.size = Pt(8.9)
        run = p.add_run(personal["location"])
        run.font.size = Pt(8.9)
        run.font.color.rgb = gray_rgb
        run.font.name = "Arial"

    def add_section_header(title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(12.7)
        run.font.color.rgb = black_rgb
        run.font.name = "Georgia"
        # Add a thin line (border below)
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "000000",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Professional Summary
    add_section_header("Professional Summary")
    summary = (content.get("professional_summary") or "").strip()
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(summary.replace("\n", " "))
        run.font.size = Pt(9.5)
        run.font.color.rgb = gray_rgb
        run.font.name = "Arial"

    # Work Experience
    work = content.get("work_experience") or []
    if work:
        add_section_header("Work Experience")
        for role in work:
            company = role.get("company") or ""
            loc = role.get("location") or ""
            title = role.get("title") or ""
            dates = role.get("dates") or ""
            if isinstance(dates, dict):
                dates = f"{dates.get('start', '')} \u2013 {dates.get('end', '')}"

            # Company + Location
            p = doc.add_paragraph()
            run = p.add_run(company)
            run.font.size = Pt(12.7)
            run.font.color.rgb = green_rgb
            run.font.name = "Arial"
            if loc:
                run = p.add_run(f"\t{loc}")
                run.font.size = Pt(10.1)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"

            # Title + Dates (with optional role_description)
            role_desc = _fix_acronym_casing((role.get("role_description") or "").strip())
            title_display = _build_title_display(title, role_desc)
            p = doc.add_paragraph()
            run = p.add_run(title_display)
            run.font.size = Pt(10.1)
            run.font.color.rgb = black_rgb
            run.font.name = "Arial"
            if dates:
                run = p.add_run(f"\t{dates}")
                run.font.size = Pt(10.1)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"

            # Bullets (justified alignment)
            for bullet in role.get("bullets") or []:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(bullet)
                run.font.size = Pt(9.5)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"

    # Key Projects
    key_projects = content.get("key_projects") or []
    if key_projects:
        add_section_header("Key Projects")
        for proj in key_projects:
            proj_name = (proj.get("name") or "").strip()
            proj_desc = (proj.get("description") or "").strip()

            # Project name (green, same as company)
            p = doc.add_paragraph()
            run = p.add_run(proj_name)
            run.font.size = Pt(12.7)
            run.font.color.rgb = green_rgb
            run.font.name = "Arial"

            # Description
            if proj_desc:
                p = doc.add_paragraph()
                run = p.add_run(proj_desc)
                run.font.size = Pt(10.1)
                run.font.color.rgb = black_rgb
                run.font.name = "Arial"

            # Bullets
            for bullet in proj.get("bullets") or []:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(bullet)
                run.font.size = Pt(9.5)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"

    # Skills
    skills = content.get("skills") or {}
    all_skills_items = []
    for cat in ("technical", "methodologies", "domains"):
        all_skills_items.extend(skills.get(cat) or [])
    if all_skills_items:
        add_section_header("Skills")
        for cat_name, cat_key in [("Technical", "technical"), ("Methodologies", "methodologies"), ("Domains", "domains")]:
            items = skills.get(cat_key) or []
            if items:
                p = doc.add_paragraph()
                run = p.add_run(f"{cat_name}: ")
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"
                run = p.add_run("  |  ".join(items))
                run.font.size = Pt(9.5)
                run.font.color.rgb = gray_rgb
                run.font.name = "Arial"

    # Awards
    awards = content.get("awards") or []
    if awards:
        add_section_header("Awards & Recognition")
        for award in awards:
            clean = (award or "").strip().lstrip("\u2022").lstrip("•").strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean)
            run.font.size = Pt(9.5)
            run.font.color.rgb = gray_rgb
            run.font.name = "Arial"

    # Education
    education = content.get("education") or []
    if education:
        add_section_header("Education")
        for edu in education:
            institution = edu.get("institution") or ""
            p = doc.add_paragraph()
            run = p.add_run(institution)
            run.font.size = Pt(12.7)
            run.font.color.rgb = green_rgb
            run.font.name = "Arial"
            edu_loc = _dedup_edu_location(institution, edu.get("location") or "")
            if edu_loc:
                run = p.add_run(f"\t{edu_loc}")
                run.font.size = Pt(10.1)
                run.font.color.rgb = gray_rgb

            p = doc.add_paragraph()
            run = p.add_run(_build_degree_with_field(edu))
            run.font.size = Pt(10.1)
            run.font.color.rgb = gray_rgb
            run.font.name = "Arial"
            dates = edu.get("dates") or ""
            if isinstance(dates, dict):
                dates = f"{dates.get('start', '')} \u2013 {dates.get('end', '')}"
            if dates:
                run = p.add_run(f"\t{dates}")
                run.font.size = Pt(10.1)
                run.font.color.rgb = gray_rgb

    # Certifications
    certs = content.get("certifications") or []
    if certs:
        add_section_header("Certifications")
        # Render all certs on a single line separated by pipes (RULE 11)
        cert_texts = [_format_cert(c) for c in certs]
        joined = "  |  ".join(cert_texts)
        p = doc.add_paragraph()
        run = p.add_run(joined)
        run.font.size = Pt(10.1)
        run.font.color.rgb = gray_rgb
        run.font.name = "Arial"

    doc.save(output_path)
    logger.info("DOCX generated: %s", output_path)


def _generate_interview_prep(reframing_log: list, content: dict) -> str:
    """Generate interview_prep.md content."""
    lines = ["# Interview Preparation Guide\n"]
    lines.append("Use this guide to prepare for questions about your resume.\n")
    lines.append("For each significantly reframed bullet, understand what actually happened ")
    lines.append("and how to explain it naturally in conversation.\n\n---\n")

    work = content.get("work_experience") or []
    role_map = {}
    for role in work:
        company = role.get("company") or "Unknown"
        title = role.get("title") or ""
        for bullet in role.get("bullets") or []:
            role_map[bullet[:60]] = {"company": company, "title": title}

    if not reframing_log:
        lines.append("\nNo significant reframing was applied. ")
        lines.append("All bullets closely match your original experience descriptions.\n")
        return "\n".join(lines)

    for entry in reframing_log:
        original = entry.get("original") or entry.get("original_pkb_text") or ""
        reframed = entry.get("reframed") or entry.get("reframed_text") or ""
        keywords_used = entry.get("jd_keywords_used") or []
        change_reason = entry.get("what_changed") or entry.get("change_reason") or ""
        interview_note = entry.get("interview_prep") or ""

        # Find which role this belongs to
        role_info = None
        for key, val in role_map.items():
            if reframed and reframed[:40] in key or key in reframed[:60]:
                role_info = val
                break

        company = role_info["company"] if role_info else "Unknown"
        title = role_info["title"] if role_info else ""

        lines.append(f"\n## {company} \u2014 {title}\n")
        lines.append(f"**Resume says:** \"{reframed}\"\n")
        lines.append(f"**What actually happened:** {original}\n")
        if change_reason:
            lines.append(f"**What changed:** {change_reason}\n")
        if keywords_used:
            lines.append(f"**JD keywords matched:** {', '.join(keywords_used)}\n")
        if interview_note:
            lines.append(f"**How to explain in interview:** {interview_note}\n")
        lines.append("---\n")

    return "\n".join(lines)


def _run_ats_parseability_check(pdf_path: str) -> dict:
    """Extract text from generated PDF and verify parseability."""
    result = {
        "pdf_path": pdf_path,
        "text_extractable": False,
        "total_chars": 0,
        "sections_found": [],
        "issues": [],
    }
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            full_text = ""
            for page in pdf.pages:
                text = page.extract_text() or ""
                full_text += text + "\n"

            result["text_extractable"] = len(full_text.strip()) > 50
            result["total_chars"] = len(full_text)

            # Check for expected sections
            expected = ["Professional Summary", "Work Experience", "Skills", "Education"]
            for section in expected:
                if section.lower() in full_text.lower():
                    result["sections_found"].append(section)
                else:
                    result["issues"].append(f"Section '{section}' not found in extracted text")

            # Check for garbled characters
            garbled_count = sum(1 for c in full_text if ord(c) > 65535)
            if garbled_count > 5:
                result["issues"].append(f"Found {garbled_count} potentially garbled characters")

    except ImportError:
        result["issues"].append("pdfplumber not installed — cannot verify parseability")
    except Exception as e:
        result["issues"].append(f"PDF parse error: {str(e)}")

    return result


def generate_output(
    formatted_content: dict,
    jd_analysis: dict,
    score_report: dict,
    keyword_report: dict = None,
    reframing_log: list = None,
    format_validation: dict = None,
    iteration_log: dict = None,
    company_name: str = None,
    candidate_name: str = None,
    pkb: dict = None,
    output_dir: str = "output",
    edit_record: dict = None,
    existing_out_folder: str = None,
    output_suffix: str = None,
    research_brief: dict = None,
) -> str:
    """Generate the full resume package (8 artifacts).

    Args:
        formatted_content: Validated resume content from formatter
        jd_analysis: Parsed JD analysis
        score_report: Full 10-component scorer v3 report
        keyword_report: Keyword coverage data
        reframing_log: List of reframing decisions
        format_validation: Validation report from formatter
        iteration_log: Scoring iteration history
        company_name: Target company name
        candidate_name: Candidate full name
        pkb: Profile Knowledge Base (for personal info)
        output_dir: Base output directory
        edit_record: Optional pre-generation edit record (content_before, content_after); write pre_generation_edit.json
        existing_out_folder: If set, use this folder instead of creating one from company_slug + date
        output_suffix: Optional unique suffix (e.g. job_id[:8]) for folder name. If None, uses HHmmss.

    Returns:
        Path to the output folder
    """
    # Resolve company name
    if not company_name:
        company_name = (jd_analysis.get("company") or "Company").strip()
    company_slug = re.sub(r"[^\w]+", "_", company_name).strip("_")

    # Resolve candidate name
    if pkb and not candidate_name:
        candidate_name = (pkb.get("personal_info") or {}).get("name") or "Candidate"
    candidate_name = candidate_name or "Candidate"
    name_slug = candidate_name.replace(" ", "_")

    if not pkb:
        # Try to load PKB from default path
        pkb_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "pkb.json")
        if os.path.exists(pkb_path):
            with open(pkb_path, "r") as f:
                pkb = json.load(f)
        else:
            pkb = {"personal_info": {"name": candidate_name}}

    # Safety net: fix pre-2023 LLM/GPT terms in experience section before quality gate
    from engine.reframer import _fix_pre_2023_tech_full
    formatted_content = _fix_pre_2023_tech_full(formatted_content, pkb)

    # Quality gate: only block on fabrication and anachronism (never block on score)
    from engine.scorer import _get_anti_pattern_issues
    anti_pattern_issues = _get_anti_pattern_issues(formatted_content, pkb)
    blocked_failures = []
    if "title_fabrication" in anti_pattern_issues:
        blocked_failures.append("title_fabrication")
    if "pre_2023_anachronistic_tech" in anti_pattern_issues:
        blocked_failures.append("pre_2023_anachronistic_tech")
    if blocked_failures:
        msg = f"Quality gate blocked: critical rule failures ({', '.join(blocked_failures)})"
        if "pre_2023_anachronistic_tech" in blocked_failures:
            from engine.scorer import _get_pre_2023_offending_role
            company, bullet_idx, term = _get_pre_2023_offending_role(formatted_content, pkb)
            if company:
                msg += f". Pre-2023 role '{company}' bullet {bullet_idx} still contains '{term}' — fix in reframer."
        raise QualityGateBlockedError(
            msg,
            blocked_reason="critical_rule_failures",
            rule13_failures=blocked_failures,
        )

    # Use existing folder (e.g. from --review flow) or create new
    if existing_out_folder and os.path.isdir(existing_out_folder):
        out_folder = existing_out_folder
    else:
        date_str = datetime.now().strftime("%Y-%m-%d")
        suffix = output_suffix if output_suffix else datetime.now().strftime("%H%M%S")
        out_folder = os.path.join(output_dir, f"{company_slug}_{date_str}_{suffix}")
        os.makedirs(out_folder, exist_ok=True)

    # --- 1. PDF ---
    pdf_filename = f"{name_slug}_{company_slug}.pdf"
    pdf_path = os.path.join(out_folder, pdf_filename)
    try:
        _generate_pdf(formatted_content, pkb, pdf_path)
    except Exception as e:
        logger.error("PDF generation failed: %s", e)
        raise

    # --- 2. DOCX ---
    docx_filename = f"{name_slug}_{company_slug}.docx"
    docx_path = os.path.join(out_folder, docx_filename)
    try:
        _generate_docx(formatted_content, pkb, docx_path)
    except Exception as e:
        logger.error("DOCX generation failed: %s", e)
        # Non-fatal — PDF is primary

    # --- 3. score_report.json ---
    with open(os.path.join(out_folder, "score_report.json"), "w") as f:
        json.dump(score_report, f, indent=2)

    # --- 4. keyword_coverage.json ---
    kw_data = keyword_report or {}
    with open(os.path.join(out_folder, "keyword_coverage.json"), "w") as f:
        json.dump(kw_data, f, indent=2)

    # --- 5. reframing_log.json ---
    rf_data = reframing_log or []
    with open(os.path.join(out_folder, "reframing_log.json"), "w") as f:
        json.dump(rf_data, f, indent=2)

    # --- 6. interview_prep.md ---
    interview_md = _generate_interview_prep(rf_data, formatted_content)
    with open(os.path.join(out_folder, "interview_prep.md"), "w") as f:
        f.write(interview_md)

    # --- 7. iteration_log.json ---
    iter_data = iteration_log or {"iterations_used": 1, "feedback_applied": [], "note": "Single pass"}
    with open(os.path.join(out_folder, "iteration_log.json"), "w") as f:
        json.dump(iter_data, f, indent=2)

    # --- 8. format_warnings.json ---
    fw_data = format_validation or {"warnings": [], "errors": []}
    # Add ATS parseability check
    ats_check = _run_ats_parseability_check(pdf_path)
    fw_data["ats_parseability"] = ats_check
    with open(os.path.join(out_folder, "format_warnings.json"), "w") as f:
        json.dump(fw_data, f, indent=2)

    # --- 9. research_brief.json (when company research was run) ---
    if research_brief:
        with open(os.path.join(out_folder, "research_brief.json"), "w") as f:
            json.dump(research_brief, f, indent=2, default=str)
        logger.info("Research brief saved to research_brief.json")

    # --- 10. pre_generation_edit.json (when user edited before PDF) ---
    if edit_record:
        pre_edit_path = os.path.join(out_folder, "pre_generation_edit.json")
        with open(pre_edit_path, "w", encoding="utf-8") as f:
            json.dump(edit_record, f, indent=2, ensure_ascii=False)
        logger.info("Pre-generation edit record saved to %s", pre_edit_path)

    logger.info("Resume package saved to: %s", out_folder)
    logger.info("  PDF: %s", pdf_filename)
    logger.info("  DOCX: %s", docx_filename)
    logger.info("  + 6 artifact files (score, keywords, reframing, interview, iteration, warnings)")

    return out_folder
